"""
DOCX Export Module

Generates a formatted Word document from a Timeline object.
Groups messages by source screenshot, showing the date, embedded image,
and speaker-attributed messages for each group.
"""

import logging
from datetime import datetime
from itertools import groupby
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.models import Message, Timeline

logger = logging.getLogger(__name__)

# Image formats that python-docx can embed
EMBEDDABLE_FORMATS = {'.jpg', '.jpeg', '.png'}

# Max image width in the document (inches)
IMAGE_WIDTH = Inches(1.25)


def export_timeline(timeline: Timeline, output_dir: Path) -> Path:
    """
    Export a Timeline to a formatted DOCX file.

    Messages are grouped by source screenshot. Each group shows the date,
    source image, and all speaker-attributed messages from that screenshot.

    Args:
        timeline: Timeline object with sorted messages
        output_dir: Directory to write the output file

    Returns:
        Path to the generated DOCX file
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y-%m-%d_%H%M%S')
    filename = f"timeline_{timestamp}.docx"
    output_path = output_dir / filename

    doc = Document()

    _add_title(doc, timeline)

    # Group consecutive messages by source file so each screenshot
    # appears once with all its messages underneath
    groups = groupby(timeline.messages, key=lambda m: m.source_file)

    for group_index, (source_file, group_messages) in enumerate(groups):
        if group_index > 0:
            _add_separator(doc)

        messages_list = list(group_messages)
        _add_screenshot_entry(doc, messages_list)

    doc.save(str(output_path))
    logger.info(f"Timeline exported to: {output_path}")

    return output_path


def _add_title(doc: Document, timeline: Timeline) -> None:
    """Add the document title and date range header."""
    title = doc.add_heading('Text Message Timeline', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date range
    dated = [m for m in timeline.messages if m.date is not None]
    if dated:
        first = dated[0].date.strftime('%B %d, %Y')
        last = dated[-1].date.strftime('%B %d, %Y')
        if first == last:
            range_text = f"Date: {first}"
        else:
            range_text = f"Date Range: {first} to {last}"
        range_para = doc.add_paragraph(range_text)
        range_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = datetime.now().strftime('%B %d, %Y at %I:%M %p')
    gen_para = doc.add_paragraph(f"Generated: {generated}")
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_para.runs[0].font.size = Pt(9)
    gen_para.runs[0].font.italic = True


def _add_separator(doc: Document) -> None:
    """Add a horizontal rule between screenshot entries."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('_' * 60)
    run.font.color.rgb = None  # Use default color


def _add_screenshot_entry(doc: Document, messages: list[Message]) -> None:
    """
    Add a group of messages from one screenshot to the document.

    All messages share the same source_file and date. The entry includes:
    - Date heading
    - Source file reference
    - Embedded source image (if supported format)
    - Each message with speaker label
    """
    first = messages[0]

    # Date heading
    if first.date:
        date_text = first.date.strftime('%B %d, %Y  %I:%M %p')
    else:
        date_text = "Date Unknown"

    doc.add_heading(date_text, level=2)

    # Source file reference
    source_para = doc.add_paragraph()
    source_run = source_para.add_run(f"Source: {first.source_file}")
    source_run.font.size = Pt(9)
    source_run.font.italic = True

    # Embed source image if available and supported
    if first.source_path and first.source_path.exists():
        ext = first.source_path.suffix.lower()
        if ext in EMBEDDABLE_FORMATS:
            try:
                doc.add_picture(str(first.source_path), width=IMAGE_WIDTH)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Could not embed image {first.source_file}: {e}")
        else:
            note = doc.add_paragraph(f"[Source file: {first.source_file} "
                                     f"({ext} format — cannot embed in document)]")
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.italic = True

    # Messages with speaker labels
    doc.add_paragraph()  # Spacing
    for message in messages:
        para = doc.add_paragraph()
        if message.speaker:
            # Bold speaker name, then regular text
            speaker_run = para.add_run(f"{message.speaker}: ")
            speaker_run.bold = True
            para.add_run(message.text)
        else:
            # No speaker attribution (fallback path)
            para.add_run(message.text)
