"""
Tests for the DOCX export module.

Tests document structure, message grouping by source file,
and speaker label formatting.
"""

from datetime import datetime
from pathlib import Path

from docx import Document

from src.models import Message, Timeline
from src.docx_exporter import export_timeline, _add_screenshot_entry


class TestExportTimeline:
    """Tests for the main export_timeline function."""

    def test_creates_docx_file(self, tmp_path):
        """Should create a .docx file in the output directory."""
        timeline = Timeline(messages=[
            Message(
                text="Hello",
                source_file="a.png",
                date=datetime(2024, 3, 15),
                speaker="Jenna",
            ),
        ])
        output_path = export_timeline(timeline, tmp_path)

        assert output_path.exists()
        assert output_path.suffix == ".docx"

    def test_contains_title(self, tmp_path):
        """The document should contain the title heading."""
        timeline = Timeline(messages=[
            Message(
                text="Hi",
                source_file="a.png",
                date=datetime(2024, 3, 15),
                speaker="Jenna",
            ),
        ])
        output_path = export_timeline(timeline, tmp_path)

        doc = Document(str(output_path))
        all_text = [p.text for p in doc.paragraphs]
        assert any("Text Message Timeline" in t for t in all_text)

    def test_contains_date_heading(self, tmp_path):
        """The document should contain the message date as a heading."""
        timeline = Timeline(messages=[
            Message(
                text="Hi",
                source_file="a.png",
                date=datetime(2024, 3, 15, 14, 30),
                speaker="Jenna",
            ),
        ])
        output_path = export_timeline(timeline, tmp_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "March 15, 2024" in all_text

    def test_empty_timeline(self, tmp_path):
        """Empty timeline should still produce a valid document."""
        timeline = Timeline(messages=[])
        output_path = export_timeline(timeline, tmp_path)

        assert output_path.exists()
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0  # At least the title


class TestAddScreenshotEntry:
    """Tests for _add_screenshot_entry (message formatting)."""

    def test_speaker_name_is_bold(self):
        """Speaker name should be rendered as a bold run."""
        doc = Document()
        messages = [
            Message(
                text="Hello there",
                source_file="chat.png",
                date=datetime(2024, 3, 15),
                speaker="Jenna",
            ),
        ]
        _add_screenshot_entry(doc, messages)

        # Find the paragraph with the speaker name
        found = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text == "Jenna: " and run.bold:
                    found = True
        assert found, "Speaker name should be bold"

    def test_message_text_follows_speaker(self):
        """Message text should appear after the speaker name."""
        doc = Document()
        messages = [
            Message(
                text="Hello there",
                source_file="chat.png",
                date=datetime(2024, 3, 15),
                speaker="Jenna",
            ),
        ]
        _add_screenshot_entry(doc, messages)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jenna: Hello there" in all_text

    def test_multiple_messages_all_appear(self):
        """All messages from a screenshot should appear in the document."""
        doc = Document()
        messages = [
            Message(text="Hey", source_file="a.png", date=datetime(2024, 3, 15), speaker="Jenna"),
            Message(text="Hi!", source_file="a.png", date=datetime(2024, 3, 15), speaker="Mike"),
            Message(text="How are you?", source_file="a.png", date=datetime(2024, 3, 15), speaker="Jenna"),
        ]
        _add_screenshot_entry(doc, messages)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jenna: Hey" in all_text
        assert "Mike: Hi!" in all_text
        assert "Jenna: How are you?" in all_text

    def test_no_speaker_renders_plain_text(self):
        """Messages without speaker should render without a label."""
        doc = Document()
        messages = [
            Message(
                text="Plain message",
                source_file="a.png",
                date=datetime(2024, 3, 15),
                speaker=None,
            ),
        ]
        _add_screenshot_entry(doc, messages)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Plain message" in all_text
        # Should NOT have a colon prefix
        assert ": Plain message" not in all_text

    def test_source_file_shown(self):
        """Source filename should appear in the entry."""
        doc = Document()
        messages = [
            Message(text="Hi", source_file="IMG_6019.PNG", date=datetime(2024, 3, 15), speaker="Jenna"),
        ]
        _add_screenshot_entry(doc, messages)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "IMG_6019.PNG" in all_text

    def test_undated_shows_date_unknown(self):
        """Messages without a date should show 'Date Unknown' heading."""
        doc = Document()
        messages = [
            Message(text="No date", source_file="a.png", speaker="Jenna"),
        ]
        _add_screenshot_entry(doc, messages)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Date Unknown" in all_text


class TestGrouping:
    """Tests for message grouping by source file in export."""

    def test_different_screenshots_get_separate_entries(self, tmp_path):
        """Messages from different source files should be separate entries."""
        timeline = Timeline(messages=[
            Message(text="A", source_file="first.png", date=datetime(2024, 3, 1), speaker="Jenna"),
            Message(text="B", source_file="first.png", date=datetime(2024, 3, 1), speaker="Mike"),
            Message(text="C", source_file="second.png", date=datetime(2024, 3, 2), speaker="Jenna"),
        ])
        output_path = export_timeline(timeline, tmp_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Both source files should appear
        assert "first.png" in all_text
        assert "second.png" in all_text

        # Both dates should appear
        assert "March 01, 2024" in all_text
        assert "March 02, 2024" in all_text

    def test_same_screenshot_grouped_together(self, tmp_path):
        """Messages from same source should appear under one heading."""
        timeline = Timeline(messages=[
            Message(text="Hey", source_file="chat.png", date=datetime(2024, 3, 15), speaker="Jenna"),
            Message(text="Hi", source_file="chat.png", date=datetime(2024, 3, 15), speaker="Mike"),
        ])
        output_path = export_timeline(timeline, tmp_path)

        doc = Document(str(output_path))
        # Source file should appear only once (in the source reference)
        source_mentions = sum(
            1 for p in doc.paragraphs if "Source: chat.png" in p.text
        )
        assert source_mentions == 1
